from docx import Document

def create_resume(name, contact_info, education, experience, skills):
    # Create a new Document
    doc = Document()

    # Add a heading with the name
    doc.add_heading(name, level=1)

    # Add contact information
    doc.add_paragraph("Contact Information:")
    doc.add_paragraph(contact_info)

    # Add education section
    doc.add_heading("Education", level=2)
    for degree in education:
        doc.add_paragraph(degree)

    # Add experience section
    doc.add_heading("Experience", level=2)
    for job in experience:
        doc.add_paragraph(job)

    # Add skills section
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(skills))

    # Save the document
    doc.save("resume.docx")

if __name__ == "__main__":
    # Example data
    name = "John Doe"
    contact_info = "123 Main St, City, State, 12345 | john.doe@email.com | (555) 555-5555"
    education = ["Bachelor of Science in Computer Science, University of XYZ, 2020-2024"]
    experience = ["Software Developer Intern, XYZ Company, Summer 2023"]
    skills = ["Python", "Java", "Web Development", "Problem Solving"]

    # Create the resume
    create_resume(name, contact_info, education, experience, skills)
